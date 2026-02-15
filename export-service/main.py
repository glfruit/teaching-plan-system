# Export Service - Python-docx Word Generator
# 教案导出服务 - 使用 python-docx 生成 Word 文档

from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import io
import uvicorn
from typing import Optional
from urllib.parse import quote

app = FastAPI(title="Teaching Plan Export Service")


class TeachingPlanData(BaseModel):
    """教案数据模型"""
    title: str
    courseName: str
    className: str
    duration: int
    objectives: Optional[str] = ""
    keyPoints: Optional[str] = ""
    process: Optional[str] = ""
    blackboard: Optional[str] = ""
    reflection: Optional[str] = ""
    methods: Optional[str] = ""
    resources: Optional[str] = ""
    teacherName: Optional[str] = ""
    department: Optional[str] = ""


def set_chinese_font(run, font_name='宋体', font_size=10.5, bold=False):
    """设置中文字体"""
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def html_to_docx_text(paragraph, html_content):
    """将简单 HTML 转换为 docx 段落文本"""
    import re
    from html.parser import HTMLParser
    
    class HTMLToTextParser(HTMLParser):
        def __init__(self):
            super().__init__()
            self.text_parts = []
            self.current_tag = None
            
        def handle_starttag(self, tag, attrs):
            self.current_tag = tag
            
        def handle_endtag(self, tag):
            self.current_tag = None
            
        def handle_data(self, data):
            if data.strip():
                self.text_parts.append((data, self.current_tag))
    
    # 清理 HTML 标签
    text = re.sub(r'<br\s*/?>', '\n', html_content)
    text = re.sub(r'</p>', '\n\n', text)
    text = re.sub(r'<[^>]+>', '', text)
    return text.strip()


def create_teaching_plan_docx(data: TeachingPlanData) -> io.BytesIO:
    """创建教案 Word 文档"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)
    
    # 标题
    title = doc.add_heading(level=0)
    title_run = title.add_run(data.title)
    set_chinese_font(title_run, font_name='黑体', font_size=18, bold=True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息表格
    doc.add_paragraph()
    info_table = doc.add_table(rows=2, cols=4)
    info_table.style = 'Table Grid'
    
    # 第一行
    row1 = info_table.rows[0].cells
    row1[0].text = '课程名称'
    row1[1].text = data.courseName
    row1[2].text = '授课班级'
    row1[3].text = data.className
    
    # 第二行
    row2 = info_table.rows[1].cells
    row2[0].text = '课时长度'
    row2[1].text = f"{data.duration} 分钟"
    row2[2].text = '授课教师'
    row2[3].text = data.teacherName or '待定'
    
    # 设置表格字体
    for row in info_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run, font_name='宋体', font_size=10.5)
    
    doc.add_paragraph()
    
    # 教学目标
    if data.objectives and data.objectives.strip() and data.objectives != '<p></p>':
        heading = doc.add_heading(level=1)
        run = heading.add_run('一、教学目标')
        set_chinese_font(run, font_name='黑体', font_size=14, bold=True)
        
        content = html_to_docx_text(doc.add_paragraph(), data.objectives)
        if content:
            p = doc.add_paragraph()
            run = p.add_run(content)
            set_chinese_font(run, font_name='宋体', font_size=10.5)
    
    # 重点难点
    if data.keyPoints and data.keyPoints.strip() and data.keyPoints != '<p></p>':
        heading = doc.add_heading(level=1)
        run = heading.add_run('二、重点难点')
        set_chinese_font(run, font_name='黑体', font_size=14, bold=True)
        
        content = html_to_docx_text(doc.add_paragraph(), data.keyPoints)
        if content:
            p = doc.add_paragraph()
            run = p.add_run(content)
            set_chinese_font(run, font_name='宋体', font_size=10.5)
    
    # 教学方法
    if data.methods and data.methods.strip():
        heading = doc.add_heading(level=1)
        run = heading.add_run('三、教学方法')
        set_chinese_font(run, font_name='黑体', font_size=14, bold=True)
        
        p = doc.add_paragraph()
        run = p.add_run(data.methods)
        set_chinese_font(run, font_name='宋体', font_size=10.5)
    
    # 教学资源
    if data.resources and data.resources.strip():
        heading = doc.add_heading(level=1)
        run = heading.add_run('四、教学资源')
        set_chinese_font(run, font_name='黑体', font_size=14, bold=True)
        
        p = doc.add_paragraph()
        run = p.add_run(data.resources)
        set_chinese_font(run, font_name='宋体', font_size=10.5)
    
    # 教学过程
    if data.process and data.process.strip() and data.process != '<p></p>':
        heading = doc.add_heading(level=1)
        run = heading.add_run('五、教学过程')
        set_chinese_font(run, font_name='黑体', font_size=14, bold=True)
        
        content = html_to_docx_text(doc.add_paragraph(), data.process)
        if content:
            p = doc.add_paragraph()
            run = p.add_run(content)
            set_chinese_font(run, font_name='宋体', font_size=10.5)
    
    # 板书设计
    if data.blackboard and data.blackboard.strip() and data.blackboard != '<p></p>':
        heading = doc.add_heading(level=1)
        run = heading.add_run('六、板书设计')
        set_chinese_font(run, font_name='黑体', font_size=14, bold=True)
        
        content = html_to_docx_text(doc.add_paragraph(), data.blackboard)
        if content:
            p = doc.add_paragraph()
            run = p.add_run(content)
            set_chinese_font(run, font_name='宋体', font_size=10.5)
    
    # 教学反思
    if data.reflection and data.reflection.strip() and data.reflection != '<p></p>':
        heading = doc.add_heading(level=1)
        run = heading.add_run('七、教学反思')
        set_chinese_font(run, font_name='黑体', font_size=14, bold=True)
        
        content = html_to_docx_text(doc.add_paragraph(), data.reflection)
        if content:
            p = doc.add_paragraph()
            run = p.add_run(content)
            set_chinese_font(run, font_name='宋体', font_size=10.5)
    
    # 保存到内存
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


@app.post("/export/word")
async def export_word(data: TeachingPlanData):
    """导出教案为 Word 文档"""
    try:
        buffer = create_teaching_plan_docx(data)
        from fastapi.responses import StreamingResponse
        
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{quote(f'{data.title}.docx')}"
            }
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")


@app.get("/health")
async def health_check():
    """健康检查"""
    return {"status": "ok", "service": "export-service"}


if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
